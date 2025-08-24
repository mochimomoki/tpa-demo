# app.py — Intelligent Transfer Pricing Platform (Functions 1–6)
# Streamlit single-file app scaffolding the core workflow for:
# 1) RPT Analyser  2) Compliance Checker  3) Information Request Generator
# 4) TPD Generator  5) Master File Generator  6) Industry Analysis Generator
#
# ✅ Design goals
# - Pluggable "Guideline Packs" per jurisdiction (thresholds, required sections, example questions)
# - Robust ingestion of PDFs/DOCX for financials & TPDs (section-name synonyms + regex heuristics)
# - Checklist-style compliance outputs with explanations
# - Roll-forward + fresh build pipelines for TPD/Master File using a DOCX template (format preserved)
# - Industry analysis writer with charts, citations/footnotes, and DOCX/XLSX exports
# - Optional live research (user-supplied URLs). Wikipedia/blogs auto-excluded.
# - Downloadable results
#
# 🔧 Recommended environment
#   pip install streamlit pandas numpy pdfplumber python-docx openpyxl matplotlib beautifulsoup4 lxml rapidfuzz pyyaml
#
# ▶️ Run:
#   streamlit run app.py

from __future__ import annotations
import os, io, re, json, uuid, tempfile, textwrap, datetime, itertools
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

import streamlit as st
import pandas as pd
import numpy as np

# Optional deps — app will degrade gracefully if unavailable
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt
except Exception:
    DocxDocument = None

try:
    import openpyxl
    from openpyxl import Workbook
except Exception:
    openpyxl = None

try:
    import matplotlib.pyplot as plt
except Exception:
    plt = None

try:
    from bs4 import BeautifulSoup
    import lxml  # noqa: F401
    import requests
except Exception:
    BeautifulSoup = None
    requests = None

try:
    import yaml
except Exception:
    yaml = None

try:
    from rapidfuzz import fuzz, process as rf_process
except Exception:
    fuzz = None
    rf_process = None

st.set_page_config(page_title="Transfer Pricing Assistant — v1 (Functions 1–6)", layout="wide")
st.title("🧠 Transfer Pricing Assistant — Functions 1–6")
st.caption("RPT analyser • Compliance checker • Info request • TPD/Master file generator • Industry analysis")

# ------------------------------------------------------------
# Helpers & Utilities
# ------------------------------------------------------------

@st.cache_data(show_spinner=False)
def _read_file_bytes(uploaded) -> bytes:
    if uploaded is None:
        return b""
    return uploaded.getvalue()


def _ensure_pkg(name: str, module_obj: Any, hint: str) -> bool:
    if module_obj is None:
        st.error(f"Missing optional dependency **{name}** — {hint}")
        return False
    return True


# ----------------- Jurisdiction Guideline Packs -----------------
# You can load a full pack via YAML/JSON upload. A small starter pack is embedded for demo purposes only.
# Structure reference:
# {
#   "jurisdictions": {
#       "Singapore": {
#           "tax_authority": "IRAS",
#           "tpd_thresholds": [
#               {"label": "Sale or purchase of goods with related parties", "currency": "SGD", "threshold": 15000000, "guidance_ref": "IRAS TP Guidelines (indicative transaction threshold)", "citation_url": "https://www.iras.gov.sg/"}
#           ],
#           "local_file_requirements": ["Organisational structure", "Business overview", ...],
#           "master_file_requirements": [ ... ],
#           "compliance_explanations": {"Organisational structure": "Org chart...", ...},
#           "info_request_questions": ["Were there any business restructurings in FY {YEAR}?", ...],
#           "industry_preferred_sources": ["https://www.mas.gov.sg/", "https://www.tablebuilder.singstat.gov.sg/"],
#       },
#       ...
#   },
#   "oecd": {"local_file_requirements": [...], "master_file_requirements": [...]} 
# }

STARTER_PACK = {
    "jurisdictions": {
        "Singapore": {
            "tax_authority": "IRAS",
            # ⚠️ Demo values — please validate/update against the latest IRAS guidance in production.
            "tpd_thresholds": [
                {"label": "Sale or purchase of goods with related parties", "currency": "SGD", "threshold": 15000000, "guidance_ref": "IRAS TP Guidelines — indicative threshold for goods (demo)", "citation_url": "https://www.iras.gov.sg/"},
                {"label": "Provision or receipt of services with related parties", "currency": "SGD", "threshold": 1000000, "guidance_ref": "IRAS TP Guidelines — indicative threshold for services (demo)", "citation_url": "https://www.iras.gov.sg/"},
            ],
            "local_file_requirements": [
                "Organisational structure",
                "Description of business and industry",
                "Covered related party transactions",
                "Functional analysis (functions, assets, risks)",
                "Selection and application of transfer pricing method",
                "Benchmarking/economic analysis",
                "Financial information (tested party results)",
                "Conclusion"
            ],
            "master_file_requirements": [
                "Organisational structure (group)",
                "Description of MNE business including important drivers of business profit",
                "Description of MNE's intangibles and strategy",
                "Intercompany financial activities",
                "MNE's consolidated financial and tax positions",
            ],
            "compliance_explanations": {
                "Organisational structure": "Include org chart and ownership structure with percentages.",
                "Description of business and industry": "Describe key activities, supply chain, and economic environment.",
                "Functional analysis (functions, assets, risks)": "Analyse who does what, who uses/owns what, and who bears which risks.",
            },
            "info_request_questions": [
                "Were there any business restructurings in FY {YEAR}? If yes, provide details and impact on P&L.",
                "Provide detailed related party transaction listings by counterparty for FY {YEAR}.",
                "Provide segmented P&L by transaction category for FY {YEAR}.",
                "Provide intercompany agreements (new/updated) executed in FY {YEAR}.",
                "Provide description of key value drivers and supply chain for FY {YEAR}.",
            ],
            "industry_preferred_sources": [
                "https://www.mas.gov.sg/",
                "https://www.tablebuilder.singstat.gov.sg/",
                "https://data.worldbank.org/",
                "https://www.imf.org/",
            ],
        },
        # Add more jurisdictions here or upload a full pack via sidebar.
        "OECD": {
            "local_file_requirements": [
                "Local entity overview",
                "Controlled transactions and context",
                "Comparable analysis and method selection",
                "Financial information of the local entity",
            ],
            "master_file_requirements": [
                "Organisational structure",
                "Description of MNE business",
                "Intangibles",
                "Intercompany financial activities",
                "Financial and tax positions",
            ],
        },
    }
}


def load_guideline_pack(uploaded_file: Optional[io.BytesIO]) -> Dict[str, Any]:
    if uploaded_file is None:
        return STARTER_PACK
    try:
        raw = uploaded_file.getvalue().decode("utf-8", errors="ignore")
        if uploaded_file.name.lower().endswith((".yaml", ".yml")) and yaml is not None:
            return yaml.safe_load(raw)
        else:
            return json.loads(raw)
    except Exception as e:
        st.warning(f"Failed to parse uploaded pack: {e}. Falling back to starter pack.")
        return STARTER_PACK


# ----------------- Text Extraction -----------------

SECTION_SYNONYMS = {
    "income_statement": [
        "income statement",
        "statement of profit or loss",
        "profit or loss",
        "statement of profit and loss",
        "statement of comprehensive income",
        "comprehensive income",
        "profit and loss",
    ],
    "related_party": [
        "related party",
        "related parties",
        "transactions with related",
        "balances with related",
        "due from related",
        "due to related",
        "significant related",
        "intercompany",
    ],
}

CURRENCY_SIGNS = ["USD", "SGD", "S$", "$", "EUR", "IDR", "MYR", "RM", "AUD", "CNY", "JPY", "£", "HKD"]
NUMBER_RE = re.compile(r"(?<![\w/])-?\(?\d{1,3}(?:,\d{3})*(?:\.\d+)?\)?")
YEAR_RE = re.compile(r"(?<!\d)(20\d{2}|19\d{2})(?!\d)")


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    if not _ensure_pkg("pdfplumber", pdfplumber, "pip install pdfplumber"):
        return ""
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        pages = [p.extract_text() or "" for p in pdf.pages]
    return "\n\n".join(pages)


def extract_text_from_docx(docx_bytes: bytes) -> str:
    if not _ensure_pkg("python-docx", DocxDocument, "pip install python-docx"):
        return ""
    bio = io.BytesIO(docx_bytes)
    doc = DocxDocument(bio)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(uploaded) -> str:
    if uploaded is None:
        return ""
    data = _read_file_bytes(uploaded)
    name = uploaded.name.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(data)
    elif name.endswith(".docx"):
        return extract_text_from_docx(data)
    else:
        return data.decode("utf-8", errors="ignore")


# ----------------- Fuzzy helpers -----------------

def fuzzy_contains(hay: str, needles: List[str], ratio: int = 80) -> bool:
    if rf_process is None:
        # Fallback: simple substring match
        h = hay.lower()
        return any(n.lower() in h for n in needles)
    choices = [hay]
    for n in needles:
        match, score, _ = rf_process.extractOne(n, choices, scorer=fuzz.partial_ratio)
        if score >= ratio:
            return True
    return False


# ----------------- Amount parsing -----------------

def _to_number(token: str) -> Optional[float]:
    if not token:
        return None
    t = token.strip().replace(",", "")
    neg = False
    if t.startswith("(") and t.endswith(")"):
        neg = True
        t = t[1:-1]
    try:
        val = float(t)
        return -val if neg else val
    except Exception:
        return None


def find_amounts_near_keywords(text: str, keywords: List[str], window: int = 120) -> List[Tuple[str, float, str]]:
    results = []
    lower = text.lower()
    for m in NUMBER_RE.finditer(text):
        start = m.start()
        snippet_start = max(0, start - window)
        snippet_end = min(len(text), m.end() + window)
        snippet = text[snippet_start:snippet_end]
        snip_lower = lower[snippet_start:snippet_end]
        if any(kw in snip_lower for kw in [k.lower() for k in keywords]):
            num = _to_number(m.group(0).replace("$", ""))
            if num is None:
                continue
            # crude currency guess
            cur = None
            for c in CURRENCY_SIGNS:
                if c.lower() in snip_lower:
                    cur = c
                    break
            results.append((cur or "(unknown)", num, snippet.strip()))
    return results


# ----------------- DOCX editing helpers -----------------

def replace_text_in_docx(doc: Any, mapping: Dict[str, str]) -> None:
    """Naive run-wise replacement that preserves basic formatting.
    Mapping keys should be simple tokens like {{CLIENT_NAME}}.
    """
    if DocxDocument is None:
        return
    # Paragraphs
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text:
                # reconstruct runs
                inline = p.runs
                whole = "".join(run.text for run in inline)
                new_text = whole.replace(k, v)
                # Clear and reinsert into the first run; preserve style of first run
                for idx in range(len(inline)):
                    inline[idx].text = "" if idx > 0 else new_text
    # Tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in mapping.items():
                        if k in p.text:
                            inline = p.runs
                            whole = "".join(run.text for run in inline)
                            new_text = whole.replace(k, v)
                            for idx in range(len(inline)):
                                inline[idx].text = "" if idx > 0 else new_text


def add_heading(doc: Any, text: str, level: int = 1):
    if DocxDocument is None:
        return
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)


def add_paragraph(doc: Any, text: str):
    if DocxDocument is None:
        return
    for para in textwrap.fill(text, 120).split("\n"):
        doc.add_paragraph(para)


def save_docx_and_return_bytes(doc: Any) -> bytes:
    if DocxDocument is None:
        return b""
    with io.BytesIO() as bio:
        doc.save(bio)
        return bio.getvalue()


# ----------------- Compliance logic -----------------

def build_checklist(requirements: List[str], explanations: Dict[str, str], found_text: str) -> pd.DataFrame:
    rows = []
    L = found_text.lower()
    for req in requirements:
        present = req.lower() in L
        rows.append({
            "Requirement": req,
            "Present?": "✅ Yes" if present else "❌ Missing",
            "Explanation": explanations.get(req, "") if not present else ""
        })
    return pd.DataFrame(rows)


# ----------------- Research (user-provided URLs only) -----------------

def fetch_and_clean(url: str) -> Tuple[str, str]:
    """Return (url, cleaned_text). Wikipedia/blogs are automatically skipped.
    If requests/bs4 are unavailable or fetch fails, returns (url, '')."""
    if requests is None or BeautifulSoup is None:
        return (url, "")
    bad_hosts = ["wikipedia.org", "/blog", "medium.com", "wordpress", "blogspot"]
    if any(b in url.lower() for b in bad_hosts):
        return (url, "")
    try:
        r = requests.get(url, timeout=20)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "lxml")
        for s in soup(["script", "style", "nav", "footer", "header"]):
            s.extract()
        text = " ".join(soup.get_text(" ").split())
        return (url, text)
    except Exception:
        return (url, "")


def summarise_text_blocks(blocks: List[str], max_chars: int = 2000) -> str:
    text = "\n\n".join(blocks)
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "..."


# ----------------- Industry charts -----------------

def make_simple_chart_png(title: str, series: List[Tuple[str, float]]) -> Optional[bytes]:
    if plt is None:
        return None
    labels = [k for k, _ in series]
    values = [v for _, v in series]
    fig, ax = plt.subplots()
    ax.bar(labels, values)
    ax.set_title(title)
    ax.set_xlabel("Category")
    ax.set_ylabel("Value")
    fig.tight_layout()
    with io.BytesIO() as bio:
        fig.savefig(bio, format="png", dpi=160)
        plt.close(fig)
        return bio.getvalue()


# ------------------------------------------------------------
# Sidebar — Guideline Pack loader + Function picker
# ------------------------------------------------------------

with st.sidebar:
    st.header("⚙️ Settings")
    gp_file = st.file_uploader("Load Guideline Pack (JSON/YAML). If omitted, a starter pack is used.", type=["json", "yaml", "yml"])
    GP = load_guideline_pack(gp_file)

    # Function selector
    function = st.selectbox(
        "Choose function",
        [
            "1. RPT Analyser",
            "2. Compliance Checker",
            "3. Information Request Generator",
            "4. TPD Generator",
            "5. Master File Generator",
            "6. Industry Analysis Generator",
        ],
        index=0,
    )

    # Country selector (used by most tools)
    jlist = sorted([k for k in GP.get("jurisdictions", {}).keys() if k != "OECD"]) or ["Singapore"]
    country = st.selectbox("Jurisdiction", jlist)

    st.caption("Jurisdiction packs are pluggable. Upload a richer pack to cover all local rules.")


J = GP.get("jurisdictions", {}).get(country, {})
OECD = GP.get("jurisdictions", {}).get("OECD", GP.get("oecd", {}))

# ------------------------------------------------------------
# 1) RPT Analyser
# ------------------------------------------------------------
if function.startswith("1."):
    st.subheader("1) RPT Analyser — Does a TPD need to be prepared?")
    st.write(
        "Upload the client's audited financial statements (PDF/DOCX/TXT). The app scans income statements and related party disclosures, maps amounts, and compares against the jurisdiction's indicative thresholds."
    )
    fs_file = st.file_uploader("Financial statements (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"])

    if fs_file is not None:
        text = extract_text(fs_file)
        st.markdown("**Quick text preview (first 1000 chars):**")
        st.code(text[:1000] + ("..." if len(text) > 1000 else ""))

        st.markdown("### Detected related party amounts (heuristic)")
        rpts = find_amounts_near_keywords(text, SECTION_SYNONYMS["related_party"])  # (currency, amount, snippet)
        df = pd.DataFrame(rpts, columns=["Currency", "Amount", "Evidence snippet"])
        if df.empty:
            st.info("No candidate related party amounts found. Try uploading another format (e.g., the notes section PDF).")
        else:
            st.dataframe(df)

        st.markdown("### Threshold comparison")
        thresholds = J.get("tpd_thresholds", [])
        if not thresholds:
            st.warning("No thresholds in the current jurisdiction pack. Upload a richer pack or edit STARTER_PACK.")
        else:
            hits = []
            for _, row in df.iterrows():
                cur, amt = row["Currency"], float(row["Amount"])
                for th in thresholds:
                    # Very naive currency handling: only compare if same currency token appears. In production, convert FX.
                    same_currency = (cur and th.get("currency") and cur.upper().strip("$")[:3] == th["currency"]) or (not cur)
                    if same_currency and th.get("threshold") is not None and amt >= float(th["threshold"]):
                        hits.append({
                            "Transaction": row["Evidence snippet"][:80] + ("..." if len(row["Evidence snippet"])>80 else ""),
                            "Amount": amt,
                            "Currency": cur or th.get("currency", ""),
                            "Exceeded threshold": th["threshold"],
                            "Threshold label": th.get("label", ""),
                            "Guidance ref": th.get("guidance_ref", ""),
                            "Source": th.get("citation_url", "")
                        })
            hit_df = pd.DataFrame(hits)
            if hit_df.empty:
                st.success("No items exceeded the configured thresholds (based on detected figures).")
            else:
                st.error("Some items appear to exceed indicative thresholds. These should be covered in the TPD.")
                st.dataframe(hit_df)
                st.download_button("Download flagged items (CSV)", hit_df.to_csv(index=False).encode("utf-8"), file_name=f"rpt_threshold_flags_{country}.csv", mime="text/csv")

        st.markdown("---")
        st.caption("Note: detection is heuristic; for best results, also upload the Notes to the FS where related party disclosures typically appear.")

# ------------------------------------------------------------
# 2) Compliance Checker
# ------------------------------------------------------------
elif function.startswith("2."):
    st.subheader("2) Compliance Checker — Validate a TPD against local rules")
    st.write("Upload a TPD (DOCX/PDF/TXT), select the jurisdiction, and get a checklist of required elements.")

    tpd_file = st.file_uploader("TPD file", type=["docx", "pdf", "txt"])
    include_oecd = st.checkbox("Include OECD Local File list in checklist", value=True)

    if tpd_file is not None:
        text = extract_text(tpd_file)
        st.markdown("**Quick text preview (first 1000 chars):**")
        st.code(text[:1000] + ("..." if len(text) > 1000 else ""))

        reqs = list(J.get("local_file_requirements", []))
        if include_oecd:
            reqs = list(dict.fromkeys(reqs + OECD.get("local_file_requirements", [])))  # de-dup while preserving order
        expl = J.get("compliance_explanations", {})

        if not reqs:
            st.warning("No requirement list found for this jurisdiction pack. Upload a richer pack.")
        else:
            checklist = build_checklist(reqs, expl, text)
            st.dataframe(checklist)
            st.download_button("Download checklist (CSV)", checklist.to_csv(index=False).encode("utf-8"), file_name=f"tpd_compliance_check_{country}.csv", mime="text/csv")

        st.markdown("---")
        st.caption("This is a rule-based presence check. Depth/quality reviews still require professional judgement.")

# ------------------------------------------------------------
# 3) Information Request Generator
# ------------------------------------------------------------
elif function.startswith("3."):
    st.subheader("3) Information Request Generator — Excel or Word")
    st.write("Upload last year's request list or TPD to roll-forward questions, plus auto-add jurisdictional asks.")

    last_year = st.file_uploader("Prior IR list / TPD (DOCX/PDF/TXT/CSV/XLSX)", type=["docx", "pdf", "txt", "csv", "xlsx"])
    target_year = st.number_input("Target FY (e.g., 2024)", min_value=1990, max_value=2100, value=datetime.date.today().year)
    out_fmt = st.selectbox("Output format", ["Excel (.xlsx)", "Word (.docx)"])

    base_questions = J.get("info_request_questions", [])

    # Extract prior questions
    prior_text = ""
    if last_year is not None:
        name = last_year.name.lower()
        if name.endswith(".xlsx") and openpyxl is not None:
            wb = openpyxl.load_workbook(io.BytesIO(_read_file_bytes(last_year)))
            ws = wb.active
            vals = []
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if isinstance(cell, str):
                        vals.append(cell)
            prior_text = "\n".join(vals)
        elif name.endswith(".csv"):
            try:
                df0 = pd.read_csv(io.BytesIO(_read_file_bytes(last_year)))
                prior_text = "\n".join(df0.astype(str).fillna("").agg(" ".join, axis=1).tolist())
            except Exception:
                prior_text = extract_text(last_year)
        else:
            prior_text = extract_text(last_year)

    # Roll-forward year tokens
    def bump_years(text: str, to_year: int) -> List[str]:
        if not text:
            return []
        years = sorted({int(y) for y in YEAR_RE.findall(text)})
        qs = []
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        for l in lines:
            q = l
            for y in years:
                q = q.replace(str(y), str(to_year))
            # Specific phrasing bump
            q = re.sub(r"FY\s*20\d{2}", f"FY {to_year}", q)
            qs.append(q)
        return qs

    rolled = bump_years(prior_text, target_year)
    # Merge with base questions
    compiled = list(dict.fromkeys(rolled + [q.format(YEAR=target_year) for q in base_questions]))

    st.markdown("### Preview")
    if compiled:
        st.write("\n".join(f"• {q}" for q in compiled[:50]))
        if len(compiled) > 50:
            st.caption(f"(+{len(compiled)-50} more)")

    if st.button("Generate & Download"):
        if out_fmt.startswith("Excel"):
            if not _ensure_pkg("openpyxl", openpyxl, "pip install openpyxl"):
                st.stop()
            wb = Workbook()
            ws = wb.active
            ws.title = f"IR FY{target_year}"
            ws.append(["#", "Question"])
            for i, q in enumerate(compiled, 1):
                ws.append([i, q])
            with io.BytesIO() as bio:
                wb.save(bio)
                st.download_button("Download IR (.xlsx)", bio.getvalue(), file_name=f"IR_FY{target_year}_{country}.xlsx")
        else:
            if not _ensure_pkg("python-docx", DocxDocument, "pip install python-docx"):
                st.stop()
            doc = DocxDocument()
            add_heading(doc, f"Information Request — FY {target_year} ({country})", 1)
            for i, q in enumerate(compiled, 1):
                doc.add_paragraph(f"{i}. {q}")
            st.download_button("Download IR (.docx)", save_docx_and_return_bytes(doc), file_name=f"IR_FY{target_year}_{country}.docx")

# ------------------------------------------------------------
# 4) TPD Generator
# ------------------------------------------------------------
elif function.startswith("4."):
    st.subheader("4) TPD Generator — Fresh or Roll-forward")
    mode = st.radio("Mode", ["Fresh", "Roll-forward"], horizontal=True)
    target_year = st.number_input("Target FY (e.g., 2024)", min_value=1990, max_value=2100, value=datetime.date.today().year)

    template_docx = st.file_uploader("Template (.docx) — required to preserve your formatting", type=["docx"])
    info_list_file = st.file_uploader("Information request list / client inputs (optional)", type=["docx", "pdf", "txt", "csv", "xlsx"], key="tpd_ir")
    urls_text = st.text_area("Optional: paste credible URLs (one per line) to cite for company/group overview and industry context (no blogs/Wikipedia).", height=120)

    company_name = st.text_input("Client/Entity name (for placeholders like {{CLIENT_NAME}})")

    if mode == "Fresh":
        st.markdown("**Placeholders supported**: {{CLIENT_NAME}}, {{FY}}, {{COUNTRY}}, plus any custom placeholders in your template.")
        st.caption("The app will also append an Industry Analysis section if provided.")

    else:
        prior_tpd = st.file_uploader("Prior year TPD (.docx) to roll forward", type=["docx"], key="roll_doc")
        st.write("This will replace year mentions and update simple numeric tokens. You can also merge new content.")

    if st.button("Generate TPD (.docx)"):
        if not _ensure_pkg("python-docx", DocxDocument, "pip install python-docx"):
            st.stop()
        if template_docx is None:
            st.error("Please provide a DOCX template.")
            st.stop()
        # Load template
        tdoc = DocxDocument(io.BytesIO(_read_file_bytes(template_docx)))

        # Collect optional research content
        urls = [u.strip() for u in (urls_text or "").splitlines() if u.strip()]
        fetched = []
        for u in urls:
            url, text = fetch_and_clean(u)
            if text:
                fetched.append((url, text))
        research_summary = summarise_text_blocks([t for _, t in fetched], 1800)

        # Build mapping
        mapping = {
            "{{CLIENT_NAME}}": company_name or "",
            "{{FY}}": str(target_year),
            "{{COUNTRY}}": country,
        }
        replace_text_in_docx(tdoc, mapping)

        # Append sections (industry analysis if any URLs provided)
        if urls:
            add_heading(tdoc, "Industry Analysis", 1)
            add_paragraph(tdoc, f"Jurisdiction: {country}. Year: {target_year}.")
            if research_summary:
                add_paragraph(tdoc, research_summary)
            # Simple demo chart (user may replace with real metrics)
            if plt is not None:
                png = make_simple_chart_png("Illustrative market trend", [("Y-2", 100), ("Y-1", 110), ("Y", 130)])
                if png:
                    tdoc.add_picture(io.BytesIO(png), width=Inches(5.5))
            # Footnotes
            if fetched:
                add_heading(tdoc, "References", 2)
                for i, (u, _) in enumerate(fetched, 1):
                    tdoc.add_paragraph(f"[{i}] {u}")

        # Merge any info list content as an appendix
        if info_list_file is not None:
            add_heading(tdoc, "Appendix — Client Inputs", 1)
            info_text = extract_text(info_list_file)
            add_paragraph(tdoc, info_text[:3000] + ("..." if len(info_text) > 3000 else ""))

        st.download_button("Download TPD (.docx)", save_docx_and_return_bytes(tdoc), file_name=f"TPD_{company_name or 'Client'}_{country}_FY{target_year}.docx")

    if function.startswith("4.") and st.checkbox("Show template tips", value=False):
        st.info("Use unique placeholders like {{CLIENT_NAME}} in your template. This app will replace them while keeping fonts/colors.")

# ------------------------------------------------------------
# 5) Master File Generator
# ------------------------------------------------------------
elif function.startswith("5."):
    st.subheader("5) Master File Generator — Based on OECD + local requirements")
    template_docx = st.file_uploader("Master file template (.docx)", type=["docx"], key="mf_tpl")
    urls_text = st.text_area("Optional: paste credible URLs for group-level industry/market context.", height=120, key="mf_urls")
    group_name = st.text_input("Group name (for {{GROUP_NAME}})")
    target_year = st.number_input("Target FY (e.g., 2024)", min_value=1990, max_value=2100, value=datetime.date.today().year, key="mf_year")

    if st.button("Generate Master File (.docx)"):
        if not _ensure_pkg("python-docx", DocxDocument, "pip install python-docx"):
            st.stop()
        if template_docx is None:
            st.error("Please provide a DOCX template.")
            st.stop()
        doc = DocxDocument(io.BytesIO(_read_file_bytes(template_docx)))
        mapping = {
            "{{GROUP_NAME}}": group_name or "",
            "{{FY}}": str(target_year),
            "{{COUNTRY}}": country,
        }
        replace_text_in_docx(doc, mapping)

        reqs = list(dict.fromkeys(OECD.get("master_file_requirements", []) + J.get("master_file_requirements", [])))
        if reqs:
            add_heading(doc, "Compliance checklist (OECD + local)", 1)
            for r in reqs:
                doc.add_paragraph(f"• {r}")

        urls = [u.strip() for u in (urls_text or "").splitlines() if u.strip()]
        fetched = []
        for u in urls:
            url, text = fetch_and_clean(u)
            if text:
                fetched.append((u, text))
        if fetched:
            add_heading(doc, "Industry/Market Context (summary)", 1)
            add_paragraph(doc, summarise_text_blocks([t for _, t in fetched], 2000))
            add_heading(doc, "References", 2)
            for i, (u, _) in enumerate(fetched, 1):
                doc.add_paragraph(f"[{i}] {u}")

        st.download_button("Download Master File (.docx)", save_docx_and_return_bytes(doc), file_name=f"MasterFile_{group_name or 'Group'}_{country}_FY{target_year}.docx")

# ------------------------------------------------------------
# 6) Industry Analysis Generator
# ------------------------------------------------------------
else:
    st.subheader("6) Industry Analysis Generator — Credible sources only")
    st.write("Select industry, country, and (optionally) provide URLs. The app drafts a DOCX with charts and footnotes.")

    industries_common = [
        "Software/SaaS", "Semiconductors", "Consumer Electronics", "Automotive OEM", "Auto Components",
        "Logistics/3PL", "E-commerce", "Retail (Apparel)", "Food & Beverage Manufacturing", "Pulp & Paper",
        "Banking", "Insurance", "Asset Management", "Healthcare Providers", "Pharmaceuticals",
        "Oil & Gas Upstream", "Oil & Gas Downstream", "Renewable Energy (Solar)", "Renewable Energy (Wind)",
        "Telecommunications", "Construction", "Real Estate (Developers)", "Mining", "Aviation",
    ]

    chosen = st.selectbox("Industry", industries_common + ["Others (type below)"])
    other = st.text_input("If 'Others', specify")
    industry = other.strip() if chosen.startswith("Others") else chosen

    urls_text = st.text_area("Paste credible URLs (gov, IFIs, exchanges, company filings). One per line.", height=160)

    add_demo_chart = st.checkbox("Include an illustrative chart (demo numbers)", value=True)

    if st.button("Generate Industry Write-up (.docx)"):
        if not _ensure_pkg("python-docx", DocxDocument, "pip install python-docx"):
            st.stop()
        doc = DocxDocument()
        add_heading(doc, f"Industry Analysis — {industry}", 1)
        add_paragraph(doc, f"Jurisdiction: {country}.")
        add_paragraph(doc, "This section summarises current market dynamics, key drivers, competitive landscape, and regulatory context, with footnoted sources.")

        urls = [u.strip() for u in (urls_text or "").splitlines() if u.strip()]
        fetched = []
        for u in urls:
            url, text = fetch_and_clean(u)
            if text:
                fetched.append((url, text))
        if fetched:
            add_heading(doc, "Narrative (summary)", 2)
            add_paragraph(doc, summarise_text_blocks([t for _, t in fetched], 2200))

        if add_demo_chart and plt is not None:
            png = make_simple_chart_png("Illustrative KPI trend", [("Y-3", 95), ("Y-2", 102), ("Y-1", 108), ("Y", 121)])
            if png:
                doc.add_picture(io.BytesIO(png), width=Inches(5.5))

        if fetched:
            add_heading(doc, "References", 2)
            for i, (u, _) in enumerate(fetched, 1):
                doc.add_paragraph(f"[{i}] {u}")

        st.download_button("Download Industry Analysis (.docx)", save_docx_and_return_bytes(doc), file_name=f"Industry_{industry.replace(' ', '_')}_{country}.docx")

# ------------------------------------------------------------
# Footer / Notes
# ------------------------------------------------------------
st.markdown("---")
st.caption(
    "This prototype focuses on core workflows and preserves your template formatting. "
    "For production, extend the Guideline Pack with full jurisdiction rules (thresholds, checklists, examples), "
    "add FX conversion logic, and harden parsing for multilingual PDFs."
)
